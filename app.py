import os
import json
import shutil
import requests
import re
from datetime import datetime
from flask import Flask, request, jsonify, render_template, send_file
from docx import Document

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

app = Flask(
    __name__,
    template_folder=os.path.join(BASE_DIR, "templates"),
    static_folder=os.path.join(BASE_DIR, "static"),
)

UPLOAD_FOLDER = os.path.join(BASE_DIR, "resume")
OUTPUT_FOLDER = os.path.join(BASE_DIR, "outputs")


def _get_chat_url_and_headers(provider, lm_url, api_key):
    """Return (chat_completions_url, extra_headers) for the given provider."""
    local_providers = {"lmstudio", "ollama", "custom"}
    if provider in local_providers:
        base = (lm_url or PROVIDER_BASE_URLS.get(provider, "http://localhost:1234")).rstrip("/")
    else:
        base = PROVIDER_BASE_URLS.get(provider, "").rstrip("/")
    
    headers = {"Authorization": f"Bearer {api_key}"} if api_key else {}
    if provider == "openrouter":
        headers["HTTP-Referer"] = "https://resume-tailor.onrender.com"
    return f"{base}/v1/chat/completions", headers


def _tailor_anthropic(data):
    """Handle resume tailoring via the Anthropic Messages API."""
    api_key = data.get("api_key", "").strip()
    model = data.get("model", "claude-sonnet-4-6").strip()
    jd_text = data.get("jd_text", "").strip()
    job_title = data.get("job_title", "").strip()
    max_tokens = int(data.get("max_tokens", 8192))

    if not api_key:
        return jsonify({"error": "Anthropic API key is required."}), 400
    if not jd_text:
        return jsonify({"error": "Job description cannot be empty."}), 400
    if not os.path.exists(DEFAULT_RESUME):
        return jsonify({"error": "No resume found - please upload one first."}), 400

    resume_text = extract_resume_text(DEFAULT_RESUME)
    resume_paras = get_resume_paragraphs(DEFAULT_RESUME)

    combined_message = build_tailor_message(resume_text, resume_paras, jd_text)

    try:
        resp = requests.post(
            "https://api.anthropic.com/v1/messages",
            json={"model": model, "max_tokens": max_tokens, "system": SYSTEM_PROMPT,
                  "messages": [{"role": "user", "content": combined_message}]},
            headers={"x-api-key": api_key, "anthropic-version": "2023-06-01",
                     "content-type": "application/json"},
            timeout=300,
        )
        if resp.status_code == 401:
            return jsonify({"error": "Invalid Anthropic API key. Check console.anthropic.com."}), 401
        if resp.status_code == 429:
            return jsonify({"error": "Anthropic API rate limit exceeded. Try again shortly."}), 429
        resp.raise_for_status()
        raw = resp.json()["content"][0]["text"]
    except requests.exceptions.ConnectionError:
        return jsonify({"error": "Cannot connect to Anthropic API."}), 503
    except requests.exceptions.Timeout:
        return jsonify({"error": "Anthropic API timed out."}), 504
    except Exception as e:
        return jsonify({"error": f"Anthropic API error: {e}"}), 500

    raw = _strip_think(raw)

    replacements = extract_json_array(raw)
    if replacements is not None:
        replacements = _normalize_replacement_keys(replacements)
    if replacements is None or not replacements:
        preview = raw[:1200] if raw else "(empty)"
        return jsonify({"error": "Could not extract JSON array from Anthropic response.",
                        "hint": "Model returned prose without a JSON array.",
                        "raw_preview": preview}), 500

    enriched = enrich_with_sections(replacements, resume_paras)
    try:
        name_prefix = _name_prefix(resume_paras)
        out_name = make_output_name(job_title, prefix=name_prefix)
        out_path = os.path.join(OUTPUT_FOLDER, out_name)
        apply_replacements(DEFAULT_RESUME, out_path, replacements)
    except Exception as e:
        return jsonify({"error": f"Failed to write .docx: {e}"}), 500

    return jsonify({"success": True, "filename": out_name, "docx_b64": _read_docx_b64(out_path),
                    "changes_count": len(enriched), "changes": enriched})


def extract_user_name(resume_paras):
    """Extract the user's name from the resume (first non-heading, non-contact paragraph)."""
    for p in resume_paras:
        if p.get("is_heading"):
            break
        text = p.get("text", "").strip()
        if not text:
            continue
        if any(k in text.lower() for k in ("@", "linkedin", "github", "portfolio", "phone")):
            continue
        if re.search(r"\d{3}[\-\.\s]\d{3}", text):
            continue
        return text
    return ""


def _name_prefix(resume_paras):
    """Return a slug of the user's name from the resume, or 'resume' as fallback."""
    name = extract_user_name(resume_paras)
    if name:
        return re.sub(r"[^a-z0-9]+", "-", name.lower()).strip("-")[:20]
    return "resume"


def make_output_name(job_title, prefix="resume"):
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    if job_title:
        slug = re.sub(r"[^a-z0-9]+", "-", job_title.lower().strip()).strip("-")[:30]
        return f"{prefix}_{slug}_{ts}.docx"
    return f"{prefix}_tailored_{ts}.docx"
DEFAULT_RESUME = os.path.join(UPLOAD_FOLDER, "base_resume.docx")
LM_STUDIO_BASE_URL = os.environ.get("LM_STUDIO_BASE_URL", "http://localhost:1234")
LM_STUDIO_URL = f"{LM_STUDIO_BASE_URL.rstrip('/')}/v1/chat/completions"
LM_STUDIO_MODELS_URL = f"{LM_STUDIO_BASE_URL.rstrip('/')}/v1/models"

# Base URLs for cloud/local providers
PROVIDER_BASE_URLS = {
    "lmstudio":   LM_STUDIO_BASE_URL,
    "ollama":     "http://localhost:11434",
    "openai":     "https://api.openai.com",
    "groq":       "https://api.groq.com/openai",
    "openrouter": "https://openrouter.ai/api",
    "mistral":    "https://api.mistral.ai",
}

PROVIDER_DISPLAY = {
    "lmstudio": "LM Studio", "ollama": "Ollama", "openai": "OpenAI",
    "anthropic": "Anthropic", "gemini": "Gemini", "groq": "Groq",
    "openrouter": "OpenRouter", "mistral": "Mistral", "custom": "Custom",
}

SYSTEM_PROMPT = """You are an expert ATS-optimized resume writer. You tailor resumes holistically for a target role — not just injecting keywords, but reframing the entire resume to tell a coherent story that positions the candidate as a strong fit.

YOUR APPROACH — think at the resume level, not the bullet level:
- Read the ENTIRE resume and JD together. Understand the candidate's real background and the target role.
- Decide a NARRATIVE STRATEGY: what unified story should this resume tell? Which experiences are most transferable? Which need the most reframing?
- Plan ALL changes as a cohesive set. Every rewritten bullet should contribute to one consistent narrative. Distribute keywords and themes across bullets — avoid repeating the same keyword in multiple bullets.
- Be AGGRESSIVE with low-relevance bullets. If a bullet describes experience irrelevant to the target role, don't just sprinkle in a keyword — reframe it entirely to highlight a different, more relevant aspect of that same role/project. The candidate likely did more than what one bullet captures; find the angle that serves this JD.

SCOPE — how many changes to make:
- Target 8–15 paragraph replacements. Fewer than 5 means you are under-tailoring. More than 20 means you are over-editing.
- ALWAYS rewrite the summary/objective to position the candidate squarely for the target role, reflecting the JD's seniority level, domain, and core competencies.
- ALWAYS reorder skills lines to front-load JD-critical terms (this counts as a change).
- For bullet points: rewrite bullets that are irrelevant or weakly relevant to the target role. Leave bullets that already strongly support the narrative.

INDUSTRY CONTEXT:
- Identify the company's industry sector from the JD (e.g., fintech, healthtech, e-commerce, SaaS, defense, etc.).
- Adapt language, framing, and emphasis to resonate with that sector. For example: emphasize compliance/regulation for fintech, patient outcomes for healthtech, scale/throughput for e-commerce, security clearance for defense.
- Where the candidate's experience maps to the target industry, highlight that connection explicitly. Where it doesn't map directly, frame transferable skills using the industry's vocabulary.

REWRITING AGGRESSIVENESS by relevance:
- HIGH relevance bullet (already matches JD well): Skip it or make minimal keyword polish only.
- MEDIUM relevance bullet (partially related): Rewrite to shift emphasis toward the JD-relevant angle of the same experience.
- LOW relevance bullet (unrelated to target role): FULLY REFRAME — change the focus of the bullet to highlight a different aspect of what the candidate did at that company/role that IS relevant. Do not just add a keyword to an irrelevant bullet.

ROLE TITLE REWRITING:
- If a role title on the resume does NOT match the target JD's domain, REPLACE it with a JD-relevant equivalent. The goal is for every role title to read as if the candidate was always in the target field.
- Example: Resume says "Software Engineer - AI" but JD is for a DevOps role → rewrite to "Software Engineer - Infrastructure" or "DevOps Engineer" (whichever is truthful given the bullets under that role).
- Only change titles where the mismatch is clear and significant. If a title is generic enough (e.g., "Software Engineer"), leave it alone.
- The replacement title must still be plausible given the candidate's actual work at that company.

SKILLS LINE FILTERING:
- REMOVE skills that are irrelevant to the target JD. Do not just reorder — actively drop skills that would confuse an ATS or recruiter scanning for the target role.
- Example: JD is for DevOps → remove LLM-specific skills like "LangChain, RAG, Fine-tuning, Prompt Engineering, HuggingFace". These signal a different career track and dilute the narrative.
- Keep skills that are tangentially useful or broadly applicable (e.g., Python is relevant everywhere). Only remove skills that clearly belong to a different specialization.
- After removing irrelevant skills, front-load the remaining line with the most JD-critical terms.
- NEVER repeat the same skill across multiple skill categories. Each skill/technology/tool must appear in EXACTLY ONE category line. Before finalizing skills lines, cross-check all categories and deduplicate. If a skill fits multiple categories, place it in the most specific one.

HARD CONSTRAINTS:
- One-to-one replacements only. Every "original" must be copied character-for-character from the resume. No new bullets, no deletions, no reordering.
- NO SKILL REPETITION: A skill/technology must appear in at most one skills category. Duplicating a skill across categories is a critical error.
- Never fabricate metrics, technologies, tools, or certifications the candidate hasn't listed.
- Preserve every existing number, percentage, and metric exactly.
- Output ONLY a raw JSON array — no markdown fences, no prose, no explanation.

WRITING RULES:
- Bullets: max 32 words. Open with a strong past-tense action verb.
- Banned openers: Led, Managed, Worked, Helped, Assisted, Supported, Utilized, Leveraged, Spearheaded, Responsible for.
- Banned phrases: "as measured by", "in order to", "with a focus on".
- Vary sentence structure across bullets — no two consecutive bullets should start with the same verb or follow the same [Verb] [object] [result] pattern.
- If a bullet has no real metric, do not invent one. Strengthen clarity and keyword fit instead.
- Summary: under 50 words, no first person ("I/my"), mirror the JD's seniority, domain, and industry sector.
- Skills lines: reorder within existing categories to front-load JD-critical terms. You may add JD-relevant skills the candidate plausibly has based on their experience, but never add niche certifications or tools without evidence.

EXAMPLE of aggressive reframing (LOW relevance → rewritten):
Original: "Built internal dashboards using React and D3.js for sales team KPI tracking"
JD Target: DevOps Engineer
Replacement: "Engineered automated monitoring dashboards with React and Grafana integration, enabling real-time infrastructure health tracking across 15 services"
Why: Same project context (dashboards), but reframed from "sales KPIs" to "infrastructure monitoring" — a plausible angle that serves the DevOps narrative.

EXAMPLE of what to skip (HIGH relevance):
Original: "Deployed CI/CD pipelines using GitHub Actions and Terraform, cutting release cycles from 2 weeks to 3 days"
Why skip: already contains strong keywords (CI/CD, GitHub Actions, Terraform), has a concrete metric, uses a good action verb. Rewriting would only risk making it worse."""

# Static model lists for providers without a listing API
ANTHROPIC_MODELS = [
    {"id": "claude-opus-4-6"},
    {"id": "claude-sonnet-4-6"},
    {"id": "claude-haiku-4-5-20251001"},
]

OPENAI_STATIC_MODELS = [
    {"id": "gpt-4o"}, {"id": "gpt-4o-mini"}, {"id": "gpt-4-turbo"},
    {"id": "o3-mini"}, {"id": "o4-mini"},
]

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

def build_tailor_message(resume_text, resume_paras, jd_text):
    """Build the unified user message for all tailoring providers."""
    return f"""JOB DESCRIPTION:
---
{jd_text}
---

CURRENT RESUME (full text — read this to understand the candidate's full background):
---
{resume_text}
---

RESUME PARAGRAPHS (use these for output — copy "original" values verbatim):
---
{json.dumps(resume_paras, indent=2)}
---

IMPORTANT: Your ENTIRE response must be a single JSON array. No analysis, no explanation, no numbered steps, no markdown — ONLY the JSON array.

THINK HOLISTICALLY (without outputting anything):
1. Identify the target role, industry sector, seniority level, and core competencies from the JD.
2. Read the ENTIRE resume as a whole — understand the candidate's real background, companies, roles, and strengths.
3. Decide the NARRATIVE STRATEGY: what story should this resume tell for this specific role? Which experiences are most relevant? Which are least relevant and need the most reframing?
4. Score every paragraph's relevance to the target role:
   - HIGH relevance (7-10): Leave alone or make minimal keyword polish.
   - MEDIUM relevance (4-6): Rewrite to emphasize the JD-relevant angle of the same experience.
   - LOW relevance (0-3): AGGRESSIVELY reframe — rewrite the bullet to highlight a different, more relevant aspect of what the candidate did at that company/role. Change the focus entirely if needed.
5. Plan changes AS A COHESIVE SET — distribute keywords across bullets so nothing repeats unnecessarily. Ensure the resume tells one unified story, not a patchwork of isolated edits.
6. For each bullet being rewritten, consider what the OTHER bullets already say — avoid redundancy, ensure breadth of coverage across JD requirements.

OUTPUT FORMAT — your entire response must be exactly this, nothing else:
[
  {{"original": "exact paragraph text copied verbatim from the resume", "replacement": "rewritten version"}}
]
Only include paragraphs that actually changed. If nothing changed, output: []

CRITICAL: Do NOT output any analysis, steps, reasoning, or explanation. Output ONLY the JSON array."""


def extract_json_array(text):
    """
    Very robust JSON array extractor. Handles:
    - Markdown fences (```json ... ```)
    - <think> blocks already stripped upstream, but handle any leftovers
    - Model outputting analysis/prose THEN the array
    - Trailing commentary after the array
    - Single JSON object (wraps in list)
    - Individual objects scattered in text (fallback)
    - Escaped quotes, smart quotes, common model artifacts
    """
    if not text:
        return None

    # 1. Strip all markdown fences
    text = re.sub(r"```+(?:json|JSON)?\s*", "", text)
    text = re.sub(r"```+", "", text)

    # 2. Replace smart/curly quotes with straight ones
    text = text.replace("\u201c", '"').replace("\u201d", '"')
    text = text.replace("\u2018", "'").replace("\u2019", "'")

    # 3. Strip any remaining <think> blocks
    text = re.sub(r"<think>.*?</think>", "", text, flags=re.DOTALL)
    text = text.strip()

    # 4. Try parsing the whole thing as-is
    try:
        result = json.loads(text)
        if isinstance(result, list):
            return [r for r in result if isinstance(r, dict)]
        if isinstance(result, dict):
            return [result]
    except Exception:
        pass

    # 5. Find ALL [...] blocks and try each
    bracket_starts = [i for i, c in enumerate(text) if c == "["]
    for start in bracket_starts:
        # Walk from the last ] backwards to find valid JSON
        depth = 0
        for i, c in enumerate(text[start:], start):
            if c == "[":
                depth += 1
            elif c == "]":
                depth -= 1
                if depth == 0:
                    candidate = text[start : i + 1]
                    try:
                        result = json.loads(candidate)
                        if isinstance(result, list) and len(result) > 0:
                            items = [r for r in result if isinstance(r, dict)]
                            if items:
                                return items
                    except Exception:
                        pass
                    break

    # 6. Extract individual JSON objects {original: ..., replacement: ...}
    #    This handles models that output objects one by one without wrapping array
    objects = []
    brace_starts = [i for i, c in enumerate(text) if c == "{"]
    for start in brace_starts:
        depth = 0
        for i, c in enumerate(text[start:], start):
            if c == "{":
                depth += 1
            elif c == "}":
                depth -= 1
                if depth == 0:
                    candidate = text[start : i + 1]
                    try:
                        obj = json.loads(candidate)
                        if isinstance(obj, dict) and (
                            "original" in obj or "replacement" in obj
                        ):
                            objects.append(obj)
                    except Exception:
                        pass
                    break
    if objects:
        return objects

    # 7. Last resort: regex extraction of key-value pairs
    pattern = r'["\']?original["\']?\s*:\s*["\']([^"\']+)["\'].*?["\']?replacement["\']?\s*:\s*["\']([^"\']+)["\']'
    matches = re.findall(pattern, text, re.DOTALL)
    if matches:
        return [{"original": o.strip(), "replacement": r.strip()} for o, r in matches]

    return None


def _normalize_replacement_keys(items):
    """
    Normalize key name variations across different models.
    Models may use 'rewritten', 'revised', 'new_text', 'modified', 'updated',
    'new', 'changed', 'result' instead of 'replacement'.
    Similarly 'original_text', 'old', 'old_text', 'source', 'before' instead of 'original'.
    """
    ORIGINAL_ALIASES = {"original_text", "old", "old_text", "source", "before", "existing", "current"}
    REPLACEMENT_ALIASES = {"rewritten", "revised", "new_text", "modified", "updated", "new",
                           "changed", "result", "after", "replacement_text", "tailored", "optimized"}

    normalized = []
    for item in items:
        if not isinstance(item, dict):
            continue
        out = {}
        for k, v in item.items():
            lk = k.strip().lower().replace(" ", "_")
            if lk == "original" or lk in ORIGINAL_ALIASES:
                out["original"] = v
            elif lk == "replacement" or lk in REPLACEMENT_ALIASES:
                out["replacement"] = v
            else:
                out[k] = v
        if "original" in out and "replacement" in out:
            normalized.append(out)
    return normalized if normalized else items


def extract_resume_text(path):
    doc = Document(path)
    lines = []
    for p in doc.paragraphs:
        if p.text.strip():
            lines.append(p.text.strip())
    return "\n".join(lines)


def get_resume_paragraphs(path):
    """
    Return paragraphs with their index, text, style, section, and paragraph_type.

    paragraph_type values:
      - "heading"   : section header (WORK EXPERIENCE, SKILLS, etc.)
      - "title"     : name/contact line or role+company+date line (not rewriteable)
      - "summary"   : the professional summary paragraph(s)
      - "skills"    : a skills/certifications line
      - "bullet"    : a substantive experience/project bullet (primary rewrite target)
      - "other"     : anything else short/uncategorized
    """
    doc = Document(path)
    result = []
    current_section = ""
    heading_styles = {"heading 1", "heading 2", "heading 3", "heading 4"}

    # Date/role title pattern: lines containing date ranges or role descriptors
    _date_pat = re.compile(
        r"(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec|january|february|march|"
        r"april|june|july|august|september|october|november|december|\d{4})",
        re.IGNORECASE,
    )
    # Skills section keywords
    _skills_section_names = {"skills", "certifications", "education", "cert"}

    in_summary = False
    in_skills = False

    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if not text:
            continue
        style_lower = p.style.name.lower()

        # ── Detect section headings ──
        is_heading = (
            any(h in style_lower for h in heading_styles)
            or (len(text) < 50 and text.isupper())
            or style_lower in ("title", "subtitle")
        )
        if is_heading:
            current_section = text
            sec_lower = text.lower()
            in_summary = "summary" in sec_lower
            in_skills = any(k in sec_lower for k in _skills_section_names)
            result.append(
                {
                    "index": i,
                    "text": text,
                    "style": p.style.name,
                    "section": current_section,
                    "is_heading": True,
                    "paragraph_type": "heading",
                }
            )
            continue

        # ── Classify non-heading paragraphs ──
        word_count = len(text.split())

        # Name / contact line: very short, no bullet indicator, near top
        is_contact_or_name = (
            word_count <= 12
            and not text.startswith(("•", "-", "–", "*", "◦"))
            and (
                "@" in text
                or "linkedin" in text.lower()
                or "github" in text.lower()
                or "portfolio" in text.lower()
                or "phone" in text.lower()
                or re.search(r"\d{3}[\-\.\s]\d{3}", text)
            )
        )

        # Role/company/date title line: contains a date OR is short with role-like content
        # These are lines like "DevOps Intern   June 2024 - August 2024"
        # or "Sequretek Pvt. Ltd.  June 2024 - August 2024"
        is_title_line = (
            not is_contact_or_name
            and word_count <= 15
            and _date_pat.search(text)
            and not text.startswith(("•", "-", "–", "*", "◦"))
        )

        # Project title line: looks like "ProjectName - Subtitle Month Year - Month Year"
        # These have dashes and date-like content but are longer
        is_project_title = (
            not is_contact_or_name
            and not is_title_line
            and word_count <= 25
            and _date_pat.search(text)
            and " - " in text
            and not text.startswith(("•", "-", "–", "*", "◦"))
        )

        if is_contact_or_name:
            para_type = "title"
        elif is_title_line or is_project_title:
            para_type = "title"
        elif in_summary and word_count >= 15:
            para_type = "summary"
        elif in_skills:
            para_type = "skills"
        elif word_count >= 15:
            para_type = "bullet"
        else:
            para_type = "other"

        result.append(
            {
                "index": i,
                "text": text,
                "style": p.style.name,
                "section": current_section,
                "is_heading": False,
                "paragraph_type": para_type,
            }
        )

    return result


def normalize_text(s):
    """Normalize for fuzzy matching — strips, collapses whitespace, removes zero-width chars."""
    s = s.strip()
    s = re.sub(r"[​‌‍﻿ \t]", " ", s)
    s = re.sub(r"\s+", " ", s)
    return s


def apply_replacements(original_path, output_path, replacements):
    """
    Apply text replacements to a docx while preserving all formatting.

    Matching strategy (in order):
    1. Exact match after normalize_text()
    2. Fuzzy match (ratio >= 0.85) for cases where LLM slightly altered the original text

    Writing strategy:
    - Preserve the formatting (bold, italic, font, size) of the first run
    - Set first run to full new text, clear all subsequent runs
    - This keeps bullet formatting, indentation, and style intact
    """
    import difflib

    shutil.copy2(original_path, output_path)
    doc = Document(output_path)

    # Build normalized lookup: normalized_text -> new_text
    rep_map = {}  # normalized original -> replacement
    raw_map = {}  # normalized original -> raw original (for logging)
    for r in replacements:
        if not isinstance(r, dict):
            print(f"[WARN] Skipping non-dict: {repr(r)[:80]}")
            continue
        old = r.get("original", "").strip()
        new = r.get("replacement", "").strip()
        if old and new:
            rep_map[normalize_text(old)] = new
            raw_map[normalize_text(old)] = old
        else:
            print(f"[WARN] Missing keys in: {list(r.keys())}")

    if not rep_map:
        doc.save(output_path)
        return

    matched = set()

    def write_para(para, new_text):
        """Replace paragraph text while preserving run formatting."""
        if not para.runs:
            # No runs — use XML direct manipulation
            from docx.oxml.ns import qn

            for child in list(para._p):
                if child.tag.endswith("}r"):  # w:r run element
                    para._p.remove(child)
            # Add a plain run
            from docx.oxml import OxmlElement

            run_el = OxmlElement("w:r")
            text_el = OxmlElement("w:t")
            text_el.text = new_text
            text_el.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            run_el.append(text_el)
            para._p.append(run_el)
            return
        # Detect mixed bold/non-bold formatting (e.g. skills lines:
        # "Languages: " [bold] + "Python, Java, Go" [normal]).
        # Split replacement at the same boundary to preserve formatting.
        first_bold = para.runs[0].bold
        split_idx = None
        if first_bold and len(para.runs) > 1:
            for i, run in enumerate(para.runs[1:], 1):
                if not run.bold:
                    split_idx = i
                    break

        if split_idx is not None and ":" in new_text:
            # Split new text at first colon to match bold/non-bold boundary
            colon_pos = new_text.index(":")
            bold_part = new_text[: colon_pos + 1]
            normal_part = new_text[colon_pos + 1 :]
            para.runs[0].text = bold_part
            # Put normal text in the first non-bold run
            para.runs[split_idx].text = normal_part
            # Blank out all other runs
            for i, run in enumerate(para.runs[1:], 1):
                if i != split_idx:
                    run.text = ""
        else:
            # Set first run to full new text, blank out the rest
            para.runs[0].text = new_text
            for run in para.runs[1:]:
                run.text = ""

    applied = 0
    for para in doc.paragraphs:
        norm = normalize_text(para.text)
        if not norm:
            continue

        # 1. Exact match
        if norm in rep_map and norm not in matched:
            print(f"[MATCH exact] {repr(norm[:60])}")
            write_para(para, rep_map[norm])
            matched.add(norm)
            applied += 1
            continue

        # 2. Fuzzy match — find best candidate
        best_key = None
        best_ratio = 0.0
        for key in rep_map:
            if key in matched:
                continue
            ratio = difflib.SequenceMatcher(None, norm, key).ratio()
            if ratio > best_ratio:
                best_ratio = ratio
                best_key = key

        if best_key and best_ratio >= 0.85:
            print(f"[MATCH fuzzy {best_ratio:.2f}] {repr(norm[:60])}")
            write_para(para, rep_map[best_key])
            matched.add(best_key)
            applied += 1

    print(f"[INFO] Applied {applied}/{len(rep_map)} replacements")

    # Warn about unmatched
    for key in rep_map:
        if key not in matched:
            print(f"[UNMATCHED] {repr(key[:80])}")

    doc.save(output_path)


ORIGINAL_RESUME_INFO = os.path.join(UPLOAD_FOLDER, "original_filename.txt")

def resume_info_data():
    if not os.path.exists(DEFAULT_RESUME):
        return {"exists": False}
    text = extract_resume_text(DEFAULT_RESUME)
    doc = Document(DEFAULT_RESUME)
    paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    original_fn = "base_resume.docx"
    if os.path.exists(ORIGINAL_RESUME_INFO):
        with open(ORIGINAL_RESUME_INFO, "r", encoding="utf-8") as f:
            original_fn = f.read().strip()
            
    return {
        "exists": True,
        "paragraphs": len(paras),
        "words": len(text.split()),
        "filename": "base_resume.docx",
        "original_filename": original_fn,
        "paragraphs_text": paras,
    }


def enrich_with_sections(replacements, resume_paras):
    """Add section context to each replacement dict, including role/project title."""
    import difflib

    # Build lookup: normalized text -> "SECTION > Role Title" or just "SECTION"
    current_section = ""
    current_title = ""
    para_lookup = {}
    for p in resume_paras:
        if p.get("is_heading"):
            current_section = p.get("section", "")
            current_title = ""
            continue
        if p.get("paragraph_type") == "title":
            current_title = p["text"]
        key = normalize_text(p["text"])
        label = f"{current_section} > {current_title}" if current_title else current_section
        para_lookup[key] = label

    enriched = []
    for r in replacements:
        if not isinstance(r, dict):
            continue
        orig_norm = normalize_text(r.get("original", ""))
        section = para_lookup.get(orig_norm, "")
        if not section and orig_norm:
            best_key, best_ratio = None, 0.0
            for key, sec in para_lookup.items():
                ratio = difflib.SequenceMatcher(None, orig_norm, key).ratio()
                if ratio > best_ratio:
                    best_ratio = ratio
                    best_key = key
            if best_key and best_ratio >= 0.8:
                section = para_lookup[best_key]
        enriched.append({**r, "section": section})
    return enriched


@app.route("/")
def index():
    has_default = os.path.exists(DEFAULT_RESUME)
    return render_template("index.html", has_default=has_default)


@app.route("/upload-resume", methods=["POST"])
def upload_resume():
    if "file" not in request.files:
        return jsonify({"error": "No file provided"}), 400
    f = request.files["file"]
    if not f.filename.endswith(".docx"):
        return jsonify({"error": "Only .docx files are accepted"}), 400
    f.save(DEFAULT_RESUME)
    with open(ORIGINAL_RESUME_INFO, "w", encoding="utf-8") as out:
        out.write(f.filename)
    return jsonify({"success": True, **resume_info_data()})


@app.route("/resume-info")
def resume_info():
    return jsonify(resume_info_data())


def _clean_extracted_text(text: str) -> str:
    """Remove JSON blobs, CSS rules, and other non-JD noise from extracted text."""
    cleaned = []
    for line in text.split("\n"):
        stripped = line.strip()
        if not stripped:
            cleaned.append("")
            continue
        # Skip lone braces / brackets
        if re.match(r'^[{}\[\];,\s]+$', stripped):
            continue
        # Skip bare booleans / nulls (from JSON config)
        if stripped.lower() in ("true", "false", "null"):
            continue
        # Skip lines that are clearly JSON (start with { or [ and contain "key": patterns)
        if stripped.startswith(("{", "[")) and re.search(r'"[\w-]+":\s', stripped):
            continue
        # Skip CSS: lines with !important
        if "!important" in stripped:
            continue
        # Skip CSS: lines with var(--)
        if "var(--" in stripped:
            continue
        # Skip CSS @-rules
        if re.match(r'^@(media|font-face|import|keyframes|charset)\b', stripped):
            continue
        # Skip CSS selectors: .class { or #id { or tag {  (must end with {)
        if re.match(r'^[.#\w\-\[\]>:,\s*~+]+\{\s*$', stripped):
            continue
        # Skip CSS property lines: must have ; and look like "prop: value;" with CSS-specific values
        if re.search(r'#[0-9a-fA-F]{3,8}', stripped) and stripped.endswith(";"):
            continue
        if re.match(r'^[\w-]+\s*:.*;\s*$', stripped) and re.search(r'(px|em|rem|%|solid|none|auto|inherit|transparent|rgb)', stripped):
            continue
        # Skip data URIs
        if "data:image" in stripped or "base64," in stripped:
            continue
        # Skip lines that are mostly non-alphanumeric (minified code) — only very long lines
        if len(stripped) > 100:
            alnum = sum(c.isalnum() or c.isspace() for c in stripped)
            if alnum < len(stripped) * 0.3:
                continue
        cleaned.append(line)
    result = "\n".join(cleaned)
    result = re.sub(r"\n{3,}", "\n\n", result).strip()
    return result


def _extract_text_from_html(html: str) -> str:
    """Parse HTML and return cleaned job description text."""
    from bs4 import BeautifulSoup
    soup = BeautifulSoup(html, "html.parser")

    # Remove noisy elements
    for tag in soup(["script", "style", "nav", "footer", "header", "noscript",
                     "iframe", "svg", "img", "button", "form", "aside"]):
        tag.decompose()

    # Try to find the main job content container first
    # Common selectors used by job boards (Greenhouse, Lever, LinkedIn, etc.)
    jd_selectors = [
        {"id": re.compile(r"job.?desc|job.?detail|job.?content|job.?body|posting", re.I)},
        {"class_": re.compile(r"job.?desc|job.?detail|job.?content|job.?body|posting.?body|description", re.I)},
        {"attrs": {"data-testid": re.compile(r"job|description|posting", re.I)}},
        "article",
        {"role": "main"},
        "main",
    ]
    for selector in jd_selectors:
        if isinstance(selector, str):
            container = soup.find(selector)
        else:
            container = soup.find(**selector)
        if container:
            text = container.get_text(separator="\n")
            text = _clean_extracted_text(text)
            if len(text) >= 100:
                return text[:15000]

    # Fallback: use full page text
    text = soup.get_text(separator="\n")
    text = _clean_extracted_text(text)
    return text[:15000]


def _fetch_html_playwright(url: str) -> tuple:
    """Fetch page using headless browser. Returns (html, visible_text)."""
    from playwright.sync_api import sync_playwright
    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True)
        page = browser.new_page()
        page.goto(url, wait_until="networkidle", timeout=30000)

        # Wait for common JD content selectors to appear (SPA career sites)
        _jd_wait_selectors = [
            "[class*='job-desc']", "[class*='job-detail']", "[class*='job-content']",
            "[class*='description']", "[class*='posting']",
            "[data-testid*='job']", "[data-testid*='description']",
            "[id*='job-desc']", "[id*='job-detail']",
            "article", "[role='main']", "main",
        ]
        for sel in _jd_wait_selectors:
            try:
                page.wait_for_selector(sel, timeout=3000)
                page.wait_for_timeout(1000)
                break
            except Exception:
                continue

        html = page.content()
        # Grab visible text directly from the rendered DOM — much cleaner for SPAs
        visible_text = page.evaluate("() => document.body.innerText")
        browser.close()
    return html, visible_text or ""


def _extract_company_from_url(url: str) -> str:
    """Extract company name from URL patterns used by known job boards."""
    from urllib.parse import urlparse
    hostname = urlparse(url).hostname or ""

    # Greenhouse: boards.greenhouse.io/company
    if "greenhouse.io" in hostname:
        path = urlparse(url).path.strip("/").split("/")
        if path and path[0]:
            return path[0].replace("-", " ").title()

    # Lever: jobs.lever.co/company
    if "lever.co" in hostname:
        path = urlparse(url).path.strip("/").split("/")
        if path and path[0]:
            return path[0].replace("-", " ").title()

    return ""


def _extract_company_from_text(text: str) -> str:
    """Extract company name from JD body text using common intro patterns."""
    lines = [l.strip() for l in text.split("\n") if l.strip()]
    _generic = {"the company", "the organization", "the team", "this",
                "our company", "the ideal candidate", "the role",
                "the position", "this position", "the successful candidate",
                "us", "the job", "the opportunity", "this role"}

    # Check first ~15 lines for heading-style patterns
    for line in lines[:15]:
        # "Careers at Company" / "Jobs at Company" / "Role at Company"
        m = re.match(r'^[\w\s,()/-]{3,40}\s+(?:at|@)\s+([A-Z][\w\s&.\'-]{1,40})$', line)
        if m:
            return m.group(1).strip()[:50]
        # "About Company" / "Join Company" / "Why Company"
        m = re.match(r'^(?:About|Join|Why|Work at|Working at|Careers at)\s+([A-Z][\w\s&.\'-]{1,40})$', line)
        if m and m.group(1).strip().lower() not in _generic:
            return m.group(1).strip()[:50]

    full = " ".join(lines)

    # "[Company Name] is a/an/the/one of/dedicated/committed/seeking/looking..."
    m = re.search(
        r'(?:^|\.\s+)([A-Z][A-Za-z\s&.\'-]{2,40}?)\s+(?:is\s+(?:a|an|the|one|dedicated|committed|seeking|looking|hiring|currently)|are\s+(?:a|an|the|seeking|looking|hiring))',
        full
    )
    if m:
        name = m.group(1).strip()
        if name.lower() not in _generic:
            return name[:50]

    # "About [Company Name]" in body text
    m = re.search(
        r'(?:About|Why|Join|Work at|Working at)\s+([A-Z][A-Za-z\s&.\'-]{2,40})(?:\s*[?\n!.]|$)',
        full
    )
    if m:
        name = m.group(1).strip().rstrip("?!. ")
        if name.lower() not in _generic:
            return name[:50]

    return ""


def _extract_company_name(html: str, url: str = "") -> str:
    """Extract company name from page metadata and common job board patterns."""
    from bs4 import BeautifulSoup
    soup = BeautifulSoup(html, "html.parser")

    _GENERIC_NAMES = {"linkedin", "indeed", "glassdoor", "greenhouse",
                      "lever", "workday", "ziprecruiter", "careers",
                      "career", "jobs", "job", "hiring", "apply"}

    # 1. Check <meta> tags (og:site_name, author, company)
    for meta in soup.find_all("meta"):
        prop = (meta.get("property") or meta.get("name") or "").lower()
        content = (meta.get("content") or "").strip()
        if content and prop in ("og:site_name", "author", "company"):
            if content.lower() not in _GENERIC_NAMES:
                return content

    # 2. Look for common company name containers
    company_selectors = [
        {"attrs": {"data-testid": re.compile(r"company.?name", re.I)}},
        {"class_": re.compile(r"company.?name|employer.?name|org.?name", re.I)},
    ]
    for sel in company_selectors:
        el = soup.find(**sel)
        if el:
            name = el.get_text(strip=True)
            if 1 < len(name) < 60:
                return name

    # 3. Parse <title> — common patterns: "Role at Company", "Role - Company", "Company | Role"
    title = soup.title.get_text(strip=True) if soup.title else ""
    if title:
        for pattern in [
            r'(?:at|@)\s+([A-Z][\w\s&.\'-]{1,40})',      # "Role at Company"
            r'\s[-–|]\s+([A-Z][\w\s&.\'-]{1,40})$',        # "Role - Company" or "Role | Company"
            r'^([A-Z][\w\s&.\'-]{1,40})\s+[-–|]',          # "Company - Role"
        ]:
            m = re.search(pattern, title)
            if m:
                name = m.group(1).strip().rstrip("-–| ")
                if name.lower() not in _GENERIC_NAMES:
                    return name

    # 4. Extract from JD body text ("Company is a...", "About Company", etc.)
    page_text = soup.get_text(separator="\n")
    text_company = _extract_company_from_text(page_text)
    if text_company:
        return text_company

    # 5. Try extracting from URL structure (Greenhouse, Lever — not Workday as subdomain is often parent org)
    if url:
        url_company = _extract_company_from_url(url)
        if url_company:
            return url_company

    return ""


@app.route("/extract-jd", methods=["POST"])
def extract_jd():
    """Extract job description text and company name from a URL."""
    url = (request.json or {}).get("url", "").strip()
    if not url:
        return jsonify({"error": "URL is required."}), 400

    html = ""
    text = ""

    # Try static fetch first (fast)
    try:
        resp = requests.get(url, timeout=15, headers={
            "User-Agent": "Mozilla/5.0 (compatible; ResumeBot/1.0)"
        })
        resp.raise_for_status()
        html = resp.text
        text = _extract_text_from_html(html)
    except Exception:
        pass

    # Fall back to headless browser for JS-rendered pages
    if len(text) < 50:
        try:
            html, visible_text = _fetch_html_playwright(url)
            text = _extract_text_from_html(html)
            # For SPA sites, DOM innerText is often much cleaner than parsed HTML
            if len(text) < 50 and visible_text:
                text = _clean_extracted_text(visible_text)
            # If HTML parsing got noise but innerText is cleaner, prefer innerText
            elif visible_text:
                clean_visible = _clean_extracted_text(visible_text)
                if len(clean_visible) > len(text) * 0.5 and len(clean_visible) >= 100:
                    text = clean_visible
        except Exception as e:
            return jsonify({"error": f"Failed to extract text (tried static + browser): {e}"}), 400

    if len(text) < 50:
        return jsonify({"error": "Could not extract meaningful text from this page."}), 400

    company = _extract_company_name(html, url=url)
    return jsonify({"text": text, "company": company})


@app.route("/provider-models", methods=["POST"])
def provider_models():
    """Unified model listing for all providers."""
    data = request.json or {}
    provider = data.get("provider", "lmstudio").lower().strip()
    api_key = data.get("api_key", "").strip()
    lm_url = data.get("lm_url", "").strip()

    if provider == "lmstudio" and not api_key:
        api_key = os.environ.get("LM_API_TOKEN", "")

    # --- Anthropic: static list ---
    if provider == "anthropic":
        return jsonify({"models": ANTHROPIC_MODELS, "static": True})

    # --- Gemini: use existing google-genai listing ---
    if provider == "gemini":
        if not api_key:
            return jsonify({"error": "API key required"}), 400
        try:
            from google import genai
            client = genai.Client(api_key=api_key)
            raw = list(client.models.list())
            models = []
            for m in raw:
                name = m.name.split("/")[-1] if "/" in m.name else m.name
                if not any(x in name for x in ["gemini", "learnlm"]):
                    continue
                if "embedding" in name or "vision" in name:
                    continue
                tok = getattr(m, "input_token_limit", None) or getattr(m, "inputTokenLimit", None)
                models.append({"id": name, "input_token_limit": tok})
            models.sort(key=lambda x: x["id"])
            return jsonify({"models": models})
        except Exception as e:
            return jsonify({"error": str(e)}), 500

    # --- Local providers (LM Studio, Ollama, Custom) ---
    if provider in ("lmstudio", "ollama", "custom"):
        base = (lm_url or PROVIDER_BASE_URLS.get(provider, "http://localhost:1234")).rstrip("/")
        headers = {"Authorization": f"Bearer {api_key}"} if api_key else {}
        # Ollama: try /api/tags first
        if provider == "ollama":
            try:
                r = requests.get(f"{base}/api/tags", headers=headers, timeout=4)
                r.raise_for_status()
                raw = r.json().get("models", [])
                models = [{"id": m["name"]} for m in raw]
                return jsonify({"models": models, "online": True})
            except Exception:
                pass  # fall through to OpenAI-compat endpoint
        try:
            r = requests.get(f"{base}/v1/models", headers=headers, timeout=4)
            r.raise_for_status()
            raw = r.json().get("data", [])
            models = []
            for m in raw:
                compat = m.get("compatibility_type", "")
                platform = "mac" if compat == "mlx" else ("desktop" if compat in ("gguf", "exl2") else "")
                models.append({"id": m["id"], "platform": platform,
                                "compat": compat, "quant": m.get("quantization", "")})
            return jsonify({"models": models, "online": True})
        except Exception:
            return jsonify({"models": [], "online": False})

    # --- OpenAI-compatible cloud providers ---
    base = PROVIDER_BASE_URLS.get(provider, "").rstrip("/")
    if not base:
        return jsonify({"error": f"Unknown provider: {provider}"}), 400
    headers = {"Authorization": f"Bearer {api_key}"} if api_key else {}
    if not api_key:
        return jsonify({"error": "API key required"}), 400
    try:
        r = requests.get(f"{base}/v1/models", headers=headers, timeout=8)
        if r.status_code == 401:
            return jsonify({"error": "Invalid API key"}), 401
        r.raise_for_status()
        raw = r.json().get("data", [])
        # Filter to chat-capable models
        chat_keywords = ["gpt", "claude", "llama", "mistral", "mixtral", "gemma", "qwen",
                         "deepseek", "command", "sonar", "hermes", "nous"]
        models = []
        for m in raw:
            mid = m.get("id", "")
            if any(k in mid.lower() for k in chat_keywords) or provider in ("groq", "mistral"):
                models.append({"id": mid})
        if not models:
            models = [{"id": m.get("id", "")} for m in raw]
        models.sort(key=lambda x: x["id"])
        return jsonify({"models": models})
    except requests.exceptions.ConnectionError:
        return jsonify({"error": f"Cannot connect to {PROVIDER_DISPLAY.get(provider, provider)}"}), 503
    except Exception as e:
        return jsonify({"error": str(e)}), 500


def run_with_keepalive(func, *args, **kwargs):
    import queue
    import threading
    from flask import current_app, request, Response, stream_with_context
    import json

    app = current_app._get_current_object()
    environ = dict(request.environ)
    import io
    environ["wsgi.input"] = io.BytesIO(b"")
    q = queue.Queue()

    def worker():
        with app.request_context(environ):
            try:
                res = func(*args, **kwargs)
                if isinstance(res, tuple):
                    resp = res[0]
                else:
                    resp = res
                body = resp.get_data(as_text=True) if hasattr(resp, "get_data") else json.dumps(resp)
                q.put({"type": "done", "body": body})
            except Exception as e:
                q.put({"type": "error", "body": json.dumps({"error": f"Internal Error: {str(e)}"})})

    t = threading.Thread(target=worker)
    t.start()

    def generate():
        yield " "  # Immediate whitespace to start transfer and bypass 100s TTFB
        while True:
            try:
                msg = q.get(timeout=15)
                yield msg["body"]
                break
            except queue.Empty:
                yield " "

    return Response(stream_with_context(generate()), mimetype='application/json')



@app.route("/tailor", methods=["POST"])
def tailor():
    data = request.json or {}
    return run_with_keepalive(_tailor_impl, data)

def _tailor_impl(data):
    provider = data.get("provider", "lmstudio").lower().strip()

    # Route to specialized handlers
    if provider == "gemini":
        return _tailor_gemini_impl(data)
    if provider == "anthropic":
        return _tailor_anthropic(data)

    # OpenAI-compatible path: lmstudio, ollama, openai, groq, openrouter, mistral, custom
    jd_text = data.get("jd_text", "").strip()
    model = data.get("model", "").strip()
    job_title = data.get("job_title", "").strip()
    api_key = data.get("api_key", "").strip()
    lm_url = data.get("lm_url", "").strip()

    if provider == "lmstudio" and not api_key:
        api_key = os.environ.get("LM_API_TOKEN", "")

    # Advanced config (all optional)
    temperature = float(data.get("temperature", 0.3))
    max_tokens = int(data.get("max_tokens", 8192))
    top_p = float(data.get("top_p", 0.95))
    top_k = data.get("top_k")
    repeat_penalty = float(data.get("repeat_penalty", 1.1))
    seed = int(data.get("seed", -1))
    context_length = data.get("context_length")

    if not jd_text:
        return jsonify({"error": "Job description cannot be empty."}), 400
    if not os.path.exists(DEFAULT_RESUME):
        return jsonify({"error": "No resume found - please upload one first."}), 400

    chat_url, extra_headers = _get_chat_url_and_headers(provider, lm_url, api_key)
    provider_label = PROVIDER_DISPLAY.get(provider, provider)

    resume_text = extract_resume_text(DEFAULT_RESUME)
    resume_paras = get_resume_paragraphs(DEFAULT_RESUME)

    combined_message = build_tailor_message(resume_text, resume_paras, jd_text)

    local_providers = {"lmstudio", "ollama", "custom"}

    # For local providers, suppress extended thinking (e.g. Qwen3 <think> blocks)
    # so output tokens are used for JSON, not reasoning.
    if provider in local_providers and "/no_think" not in combined_message:
        combined_message = combined_message + "\n/no_think"

    messages = [
        {"role": "system", "content": SYSTEM_PROMPT},
        {"role": "user", "content": combined_message},
    ]
    # Prefill forces the model to start outputting JSON immediately instead of
    # dumping analysis steps. Using '[\n{"original":' prevents models from
    # interpreting '[' as the start of a numbered list like '[1] ...'
    if provider in local_providers:
        messages.append({"role": "assistant", "content": '[{"original":'})

    payload = {
        "model": model,
        "messages": messages,
        "temperature": temperature,
        "max_tokens": max_tokens,
        "top_p": top_p,
        "stream": True,
    }
    # Request JSON output from providers that support it
    cloud_providers = {"openai", "groq", "openrouter", "mistral"}
    if provider in cloud_providers:
        payload["response_format"] = {"type": "json_object"}
    if provider in local_providers:
        payload["repeat_penalty"] = repeat_penalty
        if top_k and int(top_k) > 0:
            payload["top_k"] = int(top_k)
        if seed != -1:
            payload["seed"] = seed
        # Default to 32768 for local models to fit resume + JD + system prompt
        payload["context_length"] = int(context_length) if context_length else 32768

    resp = None
    try:
        resp = requests.post(chat_url, json=payload, headers=extra_headers, timeout=600, stream=True)
        resp.raise_for_status()
    except requests.exceptions.ConnectionError as e:
        print(f"[ERROR] ConnectionError connecting to {provider_label}: {e}")
        return jsonify({"error": f"Cannot connect to {provider_label}. Check the URL is reachable."}), 503
    except requests.exceptions.Timeout as e:
        print(f"[ERROR] Timeout connecting to {provider_label}: {e}")
        return jsonify({"error": f"{provider_label} timed out (10 min). Model may be too slow."}), 504
    except Exception as e:
        print(f"[ERROR] Exception connecting to {provider_label}: {e}")
        return jsonify({"error": f"{provider_label} error: {e}"}), 500

    raw = '[{"original":' if provider in local_providers else ""
    finish_reason = None
    stream_error = None
    try:
        for line in resp.iter_lines():
            if not line:
                continue
            line = line.decode("utf-8") if isinstance(line, bytes) else line
            # Skip SSE event type lines
            if line.startswith("event:"):
                continue
            if line.startswith("data: "):
                line = line[6:]
            if line.strip() == "[DONE]":
                break
            try:
                chunk = json.loads(line)
                # Detect error objects from the backend
                if "error" in chunk and "choices" not in chunk:
                    err_msg = chunk["error"]
                    if isinstance(err_msg, dict):
                        err_msg = err_msg.get("message", str(err_msg))
                    stream_error = str(err_msg)
                    print(f"[ERROR] Stream returned error: {stream_error}")
                    break
                choice = chunk["choices"][0]
                delta = choice["delta"].get("content", "")
                raw += delta
                if choice.get("finish_reason"):
                    finish_reason = choice["finish_reason"]
            except Exception as e:
                print(f"[WARN] Failed to parse stream chunk: {line} - Error: {e}")
                continue
    except Exception as e:
        print(f"[ERROR] Stream reading failed: {e}")
        return jsonify({"error": f"Error reading stream: {e}"}), 500
    finally:
        if resp:
            resp.close()

    if stream_error:
        # Surface backend errors with actionable hints
        if "context length" in stream_error.lower() or "tokens to keep" in stream_error.lower():
            return jsonify({
                "error": "Input too large for this model's context window.",
                "hint": (
                    "The resume + job description + system prompt exceeds what the model can handle. "
                    "In LM Studio, go to the model's settings and increase the context length "
                    "(32,768 or higher recommended). The Context Length setting in this app "
                    "only works if the model is loaded with enough context in LM Studio."
                ),
            }), 400
        return jsonify({"error": f"Model backend error: {stream_error}"}), 500

    print(f"[DEBUG] finish_reason={finish_reason}, raw length={len(raw)} chars")

    # Detect empty or truncated responses
    stripped_raw = raw.strip().strip("[").strip()
    if not stripped_raw:
        hint = "The model returned an empty response."
        if finish_reason == "length":
            hint += " Output was cut off by the token limit — try increasing Max Tokens."
        else:
            hint += (
                " This usually means the input exceeded the model's context window."
                " Try a model with a larger context, or increase the Context Length setting."
            )
        return jsonify({"error": hint}), 500

    raw = _strip_think(raw)
    print("[DEBUG] Raw output (first 800):", raw[:800])

    replacements = extract_json_array(raw)
    if replacements is not None:
        replacements = _normalize_replacement_keys(replacements)

    # Retry: if parsing failed, ask the model to fix its output
    if not replacements:
        print("[INFO] First parse failed, attempting repair request...")
        repair_messages = [
            {"role": "system", "content": "You are a JSON repair assistant. Output ONLY a valid JSON array, nothing else."},
            {"role": "user", "content": (
                "The following text was supposed to be a JSON array of objects with "
                '"original" and "replacement" keys, but it could not be parsed. '
                "Extract the data and return ONLY a valid JSON array. "
                "No markdown fences, no explanation.\n\n" + raw[:4000]
            )},
        ]
        repair_payload = {
            "model": model,
            "messages": repair_messages,
            "temperature": 0.0,
            "max_tokens": max_tokens,
            "stream": False,
        }
        if provider in cloud_providers:
            repair_payload["response_format"] = {"type": "json_object"}
        try:
            repair_resp = requests.post(chat_url, json=repair_payload, headers=extra_headers, timeout=120)
            repair_resp.raise_for_status()
            repair_raw = repair_resp.json()["choices"][0]["message"]["content"]
            repair_raw = _strip_think(repair_raw)
            print("[DEBUG] Repair output (first 500):", repair_raw[:500])
            replacements = extract_json_array(repair_raw)
            if replacements is not None:
                replacements = _normalize_replacement_keys(replacements)
        except Exception as e:
            print(f"[WARN] Repair request failed: {e}")

    if not replacements:
        preview = raw[:1200] if raw else "(empty — model produced no output)"
        if finish_reason == "length":
            error_msg = "Model output was truncated (hit token limit) — the JSON was cut off."
            hint_msg = (
                "Increase Max Tokens, increase Context Length, "
                "or use a model with a larger context window."
            )
        else:
            error_msg = "Could not extract a JSON array from the model response."
            hint_msg = "The model may have returned analysis/prose without a JSON array. Check the raw output below."
        return (
            jsonify(
                {
                    "error": error_msg,
                    "hint": hint_msg,
                    "raw_preview": preview,
                }
            ),
            500,
        )

    enriched = enrich_with_sections(replacements, resume_paras)

    try:
        name_prefix = _name_prefix(resume_paras)
        out_name = make_output_name(job_title, prefix=name_prefix)
        out_path = os.path.join(OUTPUT_FOLDER, out_name)
        apply_replacements(DEFAULT_RESUME, out_path, replacements)
    except Exception as e:
        return jsonify({"error": f"Failed to write .docx: {e}"}), 500

    return jsonify(
        {
            "success": True,
            "filename": out_name,
            "docx_b64": _read_docx_b64(out_path),
            "changes_count": len(enriched),
            "changes": enriched,
        }
    )


def _tailor_gemini_impl(data):
    try:
        from google import genai
        from google.genai import types as genai_types
    except ImportError:
        return (
            jsonify(
                {
                    "error": "google-genai is not installed. Run: pip install google-genai"
                }
            ),
            500,
        )

    api_key = data.get("api_key", "").strip()
    model_name = data.get("model", "gemini-3.1-pro-preview").strip()
    jd_text = data.get("jd_text", "").strip()
    job_title = data.get("job_title", "").strip()
    temperature = float(data.get("temperature", 0.3))
    max_tokens = int(data.get("max_tokens", 8192))

    top_p = float(data.get("top_p", 0.95))

    if not api_key:
        return jsonify({"error": "Gemini API key is required."}), 400
    if not jd_text:
        return jsonify({"error": "Job description cannot be empty."}), 400
    if not os.path.exists(DEFAULT_RESUME):
        return jsonify({"error": "No resume found - please upload one first."}), 400

    resume_text = extract_resume_text(DEFAULT_RESUME)
    resume_paras = get_resume_paragraphs(DEFAULT_RESUME)

    prompt = build_tailor_message(resume_text, resume_paras, jd_text)

    # Models that support thinking mode
    _thinking_models = {"gemini-2.5-pro", "gemini-2.5-flash", "gemini-3"}

    try:
        client = genai.Client(api_key=api_key)

        # Use thinking mode only for models that support it
        use_thinking = any(model_name.startswith(prefix) for prefix in _thinking_models)

        if use_thinking:
            config = genai_types.GenerateContentConfig(
                system_instruction=SYSTEM_PROMPT,
                temperature=1,  # required for thinking mode
                max_output_tokens=max_tokens,
                top_p=top_p,
                thinking_config=genai_types.ThinkingConfig(thinking_budget=8192),
            )
        else:
            config = genai_types.GenerateContentConfig(
                system_instruction=SYSTEM_PROMPT,
                temperature=temperature,
                max_output_tokens=max_tokens,
                top_p=top_p,
                response_mime_type="application/json",
            )

        response = client.models.generate_content(
            model=model_name,
            contents=prompt,
            config=config,
        )
        raw = response.text

    except Exception as e:
        err_str = str(e)
        if (
            "API_KEY_INVALID" in err_str
            or "api key" in err_str.lower()
            or "401" in err_str
        ):
            return (
                jsonify(
                    {
                        "error": "Invalid Gemini API key. Check your key at aistudio.google.com."
                    }
                ),
                401,
            )
        if "quota" in err_str.lower() or "429" in err_str:
            return (
                jsonify(
                    {
                        "error": "Gemini API quota exceeded. Try again later or use a different key."
                    }
                ),
                429,
            )
        if "not found" in err_str.lower() or "404" in err_str:
            return (
                jsonify(
                    {"error": f"Model '{model_name}' not found. Check the model name."}
                ),
                404,
            )
        return jsonify({"error": f"Gemini API error: {err_str}"}), 500

    raw = _strip_think(raw)
    print("[DEBUG][Gemini] Raw output (first 800):", raw[:800])

    replacements = extract_json_array(raw)
    if replacements is not None:
        replacements = _normalize_replacement_keys(replacements)

    # Retry: if parsing failed, ask Gemini to fix its output
    if not replacements:
        print("[INFO][Gemini] First parse failed, attempting repair request...")
        try:
            repair_config = genai_types.GenerateContentConfig(
                system_instruction="You are a JSON repair assistant. Output ONLY a valid JSON array, nothing else.",
                temperature=0.0,
                max_output_tokens=max_tokens,
                response_mime_type="application/json",
            )
            repair_prompt = (
                "The following text was supposed to be a JSON array of objects with "
                '"original" and "replacement" keys, but it could not be parsed. '
                "Extract the data and return ONLY a valid JSON array. "
                "No markdown fences, no explanation.\n\n" + raw[:4000]
            )
            repair_response = client.models.generate_content(
                model=model_name, contents=repair_prompt, config=repair_config,
            )
            repair_raw = _strip_think(repair_response.text)
            print("[DEBUG][Gemini] Repair output (first 500):", repair_raw[:500])
            replacements = extract_json_array(repair_raw)
            if replacements is not None:
                replacements = _normalize_replacement_keys(replacements)
        except Exception as e:
            print(f"[WARN][Gemini] Repair request failed: {e}")

    if not replacements:
        preview = raw[:1200] if raw else "(empty)"
        return (
            jsonify(
                {
                    "error": "Could not extract a JSON array from the Gemini response.",
                    "hint": "The model may have returned prose without a JSON array.",
                    "raw_preview": preview,
                }
            ),
            500,
        )

    enriched = enrich_with_sections(replacements, resume_paras)

    try:
        name_prefix = _name_prefix(resume_paras)
        out_name = make_output_name(job_title, prefix=name_prefix)
        out_path = os.path.join(OUTPUT_FOLDER, out_name)
        apply_replacements(DEFAULT_RESUME, out_path, replacements)
    except Exception as e:
        return jsonify({"error": f"Failed to write .docx: {e}"}), 500

    return jsonify(
        {
            "success": True,
            "filename": out_name,
            "docx_b64": _read_docx_b64(out_path),
            "changes_count": len(enriched),
            "changes": enriched,
        }
    )


# ─────────────────────────────────────────────────────────────────────────────
# PIPELINE HELPERS
# ─────────────────────────────────────────────────────────────────────────────


def _strip_think(text):
    """
    Strip <think>...</think> blocks from Qwen3/DeepSeek output.
    Handles: complete blocks, orphaned </think>, truncated blocks.
    """
    if not text:
        return text
    # Case 1: complete block — strip it, keep what follows
    if "<think>" in text and "</think>" in text:
        cleaned = re.sub(r"<think>.*?</think>", "", text, flags=re.DOTALL).strip()
        if cleaned:
            return cleaned
    # Case 2: orphaned </think> — prefill consumed the opener
    # Everything before </think> is thinking; everything after is the answer
    if "</think>" in text and "<think>" not in text:
        after = text.split("</think>", 1)[1].strip()
        return after if after else text.split("</think>", 1)[0].strip()
    # Case 3: <think> no </think> — truncated, fish out any JSON
    if "<think>" in text:
        for marker in ["{", "["]:
            idx = text.rfind(marker)
            if idx != -1:
                candidate = text[idx:].strip()
                if (marker == "{" and "}" in candidate) or (
                    marker == "[" and "]" in candidate
                ):
                    return candidate
        return ""
    # Case 4: clean
    return text.strip()


# Always return JSON, never an HTML error page
@app.errorhandler(Exception)
def handle_exception(e):
    import traceback

    traceback.print_exc()
    return jsonify({"error": str(e)}), 500


@app.route("/download/<filename>")
def download(filename):
    safe = os.path.basename(filename)
    path = os.path.join(OUTPUT_FOLDER, safe)
    if not os.path.exists(path):
        return jsonify({"error": "File not found"}), 404
    return send_file(path, as_attachment=True)


def _read_docx_b64(path):
    """Read a .docx file and return it base64-encoded (for localStorage caching)."""
    import base64
    with open(path, "rb") as f:
        return base64.b64encode(f.read()).decode("utf-8")




if __name__ == "__main__":
    port = int(os.environ.get("PORT", 8000))
    print(f"\n  Resume Tailor -> http://0.0.0.0:{port}\n")
    app.run(host="0.0.0.0", port=port)
