import json
import os
import re
import shutil
from datetime import datetime

from docx import Document

from config import DEFAULT_RESUME, ORIGINAL_RESUME_INFO


def _clean_llm_json_text(text):
    text = re.sub(r"```+(?:json|JSON)?\s*", "", text)
    text = re.sub(r"```+", "", text)
    text = text.replace("\u201c", '"').replace("\u201d", '"')
    text = text.replace("\u2018", "'").replace("\u2019", "'")
    return re.sub(r"<think>.*?</think>", "", text, flags=re.DOTALL).strip()


def _extract_json_from_brackets(text):  # NOSONAR
    for start in (index for index, char in enumerate(text) if char == "["):
        depth = 0
        for index, char in enumerate(text[start:], start):
            if char == "[":
                depth += 1
            elif char == "]":
                depth -= 1
                if depth == 0:
                    candidate = text[start : index + 1]
                    try:
                        result = json.loads(candidate)
                    except Exception:
                        break
                    if isinstance(result, list):
                        items = [item for item in result if isinstance(item, dict)]
                        if items:
                            return items
                    break
    return None


def _extract_json_from_braces(text):  # NOSONAR
    objects = []
    for start in (index for index, char in enumerate(text) if char == "{"):
        depth = 0
        for index, char in enumerate(text[start:], start):
            if char == "{":
                depth += 1
            elif char == "}":
                depth -= 1
                if depth == 0:
                    candidate = text[start : index + 1]
                    try:
                        obj = json.loads(candidate)
                    except Exception:
                        break
                    if isinstance(obj, dict) and (
                        "original" in obj or "replacement" in obj
                    ):
                        objects.append(obj)
                    break
    return objects or None


def _extract_json_from_regex(text):
    pattern = r'["\']?original["\']?\s*:\s*["\']([^"\']+)["\'].*?["\']?replacement["\']?\s*:\s*["\']([^"\']+)["\']'  # NOSONAR
    matches = re.findall(pattern, text, re.DOTALL)
    if not matches:
        return None
    return [
        {"original": original.strip(), "replacement": replacement.strip()}
        for original, replacement in matches
    ]


def _best_fuzzy_key(normalized, candidates, matched, threshold):
    import difflib

    best_key = None
    best_ratio = 0.0
    for key in candidates:
        if key in matched:
            continue
        ratio = difflib.SequenceMatcher(None, normalized, key).ratio()
        if ratio > best_ratio:
            best_ratio = ratio
            best_key = key
    if best_key and best_ratio >= threshold:
        return best_key, best_ratio
    return None, 0.0


def _write_paragraph_text(paragraph, new_text, preserve_bold_split=False):  # NOSONAR
    if not paragraph.runs:
        from docx.oxml import OxmlElement

        for child in paragraph._p:
            if child.tag.endswith("}r"):
                paragraph._p.remove(child)
        run_el = OxmlElement("w:r")
        text_el = OxmlElement("w:t")
        text_el.text = new_text
        text_el.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        run_el.append(text_el)
        paragraph._p.append(run_el)
        return

    if preserve_bold_split:
        first_bold = paragraph.runs[0].bold
        split_index = None
        if first_bold and len(paragraph.runs) > 1:
            for index, run in enumerate(paragraph.runs[1:], 1):
                if not run.bold:
                    split_index = index
                    break
        if split_index is not None and ":" in new_text:
            colon_pos = new_text.index(":")
            paragraph.runs[0].text = new_text[: colon_pos + 1]
            paragraph.runs[split_index].text = new_text[colon_pos + 1 :]
            for index, run in enumerate(paragraph.runs[1:], 1):
                if index != split_index:
                    run.text = ""
            return

    paragraph.runs[0].text = new_text
    for run in paragraph.runs[1:]:
        run.text = ""


def _is_contact_line(text):
    return bool(
        any(
            keyword in text.lower()
            for keyword in ("@", "linkedin", "github", "portfolio", "phone")
        )
        or re.search(r"\d{3}[\-\.\s]\d{3}", text)
    )


def _is_section_heading_text(text, date_pattern, section_heading_keywords):
    return (
        len(text) < 60
        and text.isupper()
        and (
            len(text.split()) >= 5
            or date_pattern.search(text)
            or any(keyword in text.lower() for keyword in section_heading_keywords)
        )
    )


def _classify_paragraph_text(
    text, in_summary, in_skills, in_certifications, date_pattern
):
    word_count = len(text.split())
    is_contact_or_name = (
        word_count <= 12
        and not text.startswith(("•", "-", "–", "*", "◦"))
        and _is_contact_line(text)
    )
    is_title_line = (
        not is_contact_or_name
        and not in_summary
        and word_count <= 15
        and date_pattern.search(text)
        and not text.startswith(("•", "-", "–", "*", "◦"))
    )
    is_project_title = (
        not is_contact_or_name
        and not is_title_line
        and not in_summary
        and word_count <= 25
        and date_pattern.search(text)
        and " - " in text
        and not text.startswith(("•", "-", "–", "*", "◦"))
    )
    if is_contact_or_name or is_title_line or is_project_title:
        return "title"
    if in_summary and word_count >= 3:
        return "summary"
    if in_certifications:
        return "certification"
    if in_skills:
        return "skills"
    if word_count >= 15:
        return "bullet"
    if not text.startswith(("•", "-", "–", "*", "◦", "▪", "▸", "○")):
        return "title"
    return "other"


def _build_replacement_map(replacements):
    replacement_map = {}
    for replacement in replacements:
        if not isinstance(replacement, dict):
            print(f"[WARN] Skipping non-dict: {repr(replacement)[:80]}")
            continue
        old_text = replacement.get("original", "").strip()
        new_text = replacement.get("replacement", "").strip()
        if old_text and new_text:
            replacement_map[normalize_text(old_text)] = new_text
        else:
            print(f"[WARN] Missing keys in: {list(replacement.keys())}")
    return replacement_map


def _apply_replacement_map(document, replacement_map, preserve_bold_split=False):
    matched = set()
    applied = 0
    for paragraph in document.paragraphs:
        normalized = normalize_text(paragraph.text)
        if not normalized:
            continue
        if normalized in replacement_map and normalized not in matched:
            print(f"[MATCH exact] {repr(normalized[:60])}")
            _write_paragraph_text(
                paragraph, replacement_map[normalized], preserve_bold_split
            )
            matched.add(normalized)
            applied += 1
            continue
        best_key, best_ratio = _best_fuzzy_key(
            normalized, replacement_map, matched, 0.85
        )
        if best_key:
            print(f"[MATCH fuzzy {best_ratio:.2f}] {repr(normalized[:60])}")
            _write_paragraph_text(
                paragraph, replacement_map[best_key], preserve_bold_split
            )
            matched.add(best_key)
            applied += 1
    return applied, matched


def _best_section_label(original_norm, paragraph_lookup):
    section = paragraph_lookup.get(original_norm, "")
    if section or not original_norm:
        return section
    best_key, best_ratio = _best_fuzzy_key(original_norm, paragraph_lookup, set(), 0.8)
    return paragraph_lookup.get(best_key, "") if best_key and best_ratio >= 0.8 else ""


def extract_user_name(resume_paras):
    """Extract the user's name from the resume (first non-heading, non-contact paragraph)."""
    for paragraph in resume_paras:
        if paragraph.get("is_heading"):
            break
        text = paragraph.get("text", "").strip()
        if not text:
            continue
        if any(
            keyword in text.lower()
            for keyword in ("@", "linkedin", "github", "portfolio", "phone")
        ):
            continue
        if re.search(r"\d{3}[\-\.\s]\d{3}", text):
            continue
        return text
    return ""


def name_prefix(resume_paras):
    """Return a slug of the user's name from the resume, or 'resume' as fallback."""
    name = extract_user_name(resume_paras)
    if name:
        return re.sub(r"[^a-z0-9]+", "-", name.lower()).strip("-")[:20]
    return "resume"


def make_output_name(job_title, prefix="resume"):
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    if job_title:
        slug = re.sub(r"[^a-z0-9]+", "-", job_title.lower().strip()).strip("-")[:30]
        return f"{prefix}_{slug}_{timestamp}.docx"
    return f"{prefix}_tailored_{timestamp}.docx"


def build_tailor_message(resume_text, resume_paras, jd_text):
    """Build the unified user message for all tailoring providers."""
    return f"""JOB DESCRIPTION:
---
{jd_text}
---

CURRENT RESUME (full text — read this to understand the candidate's full background):
---
{resume_text}
---

RESUME PARAGRAPHS (use these for output — copy "original" values verbatim):
---
{json.dumps(resume_paras, indent=2)}
---

IMPORTANT: Your ENTIRE response must be a single JSON array. No analysis, no explanation, no numbered steps, no markdown — ONLY the JSON array.

THINK HOLISTICALLY (without outputting anything):
1. Identify the target role, industry sector, seniority level, and core competencies from the JD.
2. Read the ENTIRE resume as a whole — understand the candidate's real background, companies, roles, and strengths.
3. Decide the NARRATIVE STRATEGY: what story should this resume tell for this specific role? Which experiences are most relevant? Which are least relevant and need the most reframing?
4. THE PRESERVATION RULE (CRITICAL): Identify the hard skills, specific technologies (e.g., React, Python, local LLMs), proprietary project/platform names (e.g., PERCEPT CTEM, Percept XDR), and quantitative metrics in the original bullet. You MUST retain these in your rewrite. You are changing the *angle* of the achievement, not erasing the reality of what the candidate actually built.
5. Score every paragraph's relevance to the target role:
   - HIGH relevance (7-10): Leave alone or make minimal keyword polish. 
   - MEDIUM relevance (4-6): Rewrite to emphasize the JD-relevant angle of the same experience, while strictly keeping the original technologies and project names.
   - LOW relevance (0-3): AGGRESSIVELY reframe the *impact* or *business value* to match the JD, but DO NOT delete the core technologies, tools, or specific project names used. Frame the existing tech in a way that proves the JD's required competencies.
6. Plan changes AS A COHESIVE SET — distribute keywords across bullets so nothing repeats unnecessarily. Ensure the resume tells one unified story, not a patchwork of isolated edits.
7. CROSS-BULLET AUDIT: Before writing any replacement text, hold ALL planned changes in mind at once. For every replacement ask: (a) Does this verb or technology already appear in another planned replacement? If so, swap the verb or shift the emphasis angle. (b) Do all bullets under this same company/role form a coherent story — does this replacement contradict or duplicate any sibling bullet? (c) Does this replacement, combined with all others, read as the work of one consistent candidate? Revise until the full set is internally consistent.
8. PARAGRAPH TYPE GATE: Only include paragraphs with paragraph_type "bullet", "summary", or "skills" in your output. Every paragraph with paragraph_type "title", "heading", "other", or "certification" must be skipped — these must NEVER appear in your JSON array. Certification paragraphs are immutable.
OUTPUT FORMAT - your entire response must be exactly this, nothing else:
[
  {{"role_suggestions": [{{"original_title": "verbatim title line from resume", "suggested_title": "New Role Title"}}, {{"original_title": "another verbatim title line", "suggested_title": "Another Title"}}]}},
  {{"original": "exact paragraph text copied verbatim from the resume", "replacement": "rewritten version"}},
  {{"original": "exact paragraph text copied verbatim from the resume", "replacement": "rewritten version"}}
]
Rules:
- The role_suggestions object is optional — include it only if at least one experience warrants a title change. If included it MUST be the first item.
- Only include {{"original","replacement"}} pairs for paragraphs that actually changed. The summary MUST always be included as a change.
CRITICAL: Do NOT output any analysis, steps, reasoning, or explanation. Output ONLY the JSON array."""


def extract_json_array(text):
    """
    Very robust JSON array extractor. Handles:
    - Markdown fences (```json ... ```)
    - <think> blocks already stripped upstream, but handle any leftovers
    - Model outputting analysis/prose THEN the array
    - Trailing commentary after the array
    - Single JSON object (wraps in list)
    - Individual objects scattered in text (fallback)
    - Escaped quotes, smart quotes, common model artifacts
    """
    if not text:
        return None

    text = _clean_llm_json_text(text)

    try:
        result = json.loads(text)
        if isinstance(result, list):
            return [item for item in result if isinstance(item, dict)]
        if isinstance(result, dict):
            return [result]
    except Exception:
        pass

    for parser in (
        _extract_json_from_brackets,
        _extract_json_from_braces,
        _extract_json_from_regex,
    ):
        result = parser(text)
        if result:
            return result

    return None


def normalize_replacement_keys(items):  # NOSONAR
    """
    Normalize key name variations across different models.
    Models may use 'rewritten', 'revised', 'new_text', 'modified', 'updated',
    'new', 'changed', 'result' instead of 'replacement'.
    Similarly 'original_text', 'old', 'old_text', 'source', 'before' instead of 'original'.
    """
    original_aliases = {
        "original_text",
        "old",
        "old_text",
        "source",
        "before",
        "existing",
        "current",
    }
    replacement_aliases = {
        "rewritten",
        "revised",
        "new_text",
        "modified",
        "updated",
        "new",
        "changed",
        "result",
        "after",
        "replacement_text",
        "tailored",
        "optimized",
    }

    normalized = []
    for item in items:
        if not isinstance(item, dict):
            continue
        output = {}
        for key, value in item.items():
            lowered_key = key.strip().lower().replace(" ", "_")
            if lowered_key == "original" or lowered_key in original_aliases:
                output["original"] = value
            elif lowered_key == "replacement" or lowered_key in replacement_aliases:
                output["replacement"] = value
            else:
                output[key] = value
        if "original" in output and "replacement" in output:
            normalized.append(output)
    return normalized if normalized else items


def normalize_text(text):
    """Normalize for fuzzy matching — strips, collapses whitespace, removes zero-width chars,
    and unifies quote/dash variants so LLM output matches document text."""
    text = text.strip()
    text = re.sub(r"[​‌‍﻿ \t]", " ", text)
    text = re.sub(r"\s+", " ", text)
    # Unify smart quotes → straight quotes
    text = text.replace("\u201c", '"').replace("\u201d", '"')
    text = text.replace("\u2018", "'").replace("\u2019", "'")
    # Unify dash variants → plain hyphen
    text = text.replace("\u2013", "-").replace("\u2014", "-")
    # Unify non-breaking space
    text = text.replace("\u00a0", " ")
    return text


def filter_replacements_by_type(replacements, resume_paras):  # NOSONAR
    """
    Server-side safety net: drop any replacement that targets a title, heading,
    or other non-editable paragraph (as classified by get_resume_paragraphs).
    This prevents LLM hallucinations from corrupting role names, dates, or section headers.
    """
    import difflib

    allowed = {"bullet", "summary", "skills"}
    type_lookup = {
        normalize_text(paragraph["text"]): paragraph["paragraph_type"]
        for paragraph in resume_paras
        if paragraph.get("text")
    }
    filtered = []
    for replacement in replacements:
        original_norm = normalize_text(replacement.get("original", ""))
        if not original_norm:
            continue
        paragraph_type = type_lookup.get(original_norm)
        if paragraph_type is None:
            best_key = None
            best_ratio = 0.0
            for key in type_lookup:
                ratio = difflib.SequenceMatcher(None, original_norm, key).ratio()
                if ratio > best_ratio:
                    best_ratio = ratio
                    best_key = key
            if best_key and best_ratio >= 0.88:
                paragraph_type = type_lookup[best_key]
        if paragraph_type is not None and paragraph_type not in allowed:
            print(
                f"[FILTER] Blocked replacement of {paragraph_type!r} paragraph: {repr(original_norm[:60])}"
            )
            continue
        filtered.append(replacement)
    return filtered


def extract_resume_text(path):
    document = Document(path)
    lines = []
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            lines.append(paragraph.text.strip())
    return "\n".join(lines)


def get_resume_paragraphs(path):  # NOSONAR
    """
    Return paragraphs with their index, text, style, section, and paragraph_type.

    paragraph_type values:
      - "heading"   : section header (WORK EXPERIENCE, SKILLS, etc.)
      - "title"     : name/contact line or role+company+date line (not rewriteable)
      - "summary"   : the professional summary paragraph(s)
      - "skills"    : a skills/certifications line
      - "bullet"    : a substantive experience/project bullet (primary rewrite target)
      - "other"     : anything else short/uncategorized
    """
    document = Document(path)
    result = []
    current_section = ""
    heading_styles = {"heading 1", "heading 2", "heading 3", "heading 4"}

    date_pattern = re.compile(
        r"(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec|january|february|march|"
        r"april|june|july|august|september|october|november|december|\d{4})",
        re.IGNORECASE,
    )
    skills_section_names = {
        "skills",
        "technical skills",
        "core competencies",
        "expertise",
    }
    certification_section_names = {
        "certifications",
        "certification",
        "cert",
        "licenses",
        "licence",
        "licenses & certifications",
    }
    # Keywords that mark a line as a *section* heading (not just an all-caps person name)
    section_heading_keywords = frozenset(
        {
            "experience",
            "skills",
            "education",
            "summary",
            "objective",
            "profile",
            "projects",
            "certifications",
            "achievements",
            "awards",
            "publications",
            "volunteer",
            "about",
            "background",
            "contact",
            "career",
            "work",
            "employment",
            "history",
            "technical",
            "professional",
            "activities",
            "interests",
            "languages",
            "courses",
            "training",
            "references",
            "internship",
        }
    )

    in_summary = (
        True  # treat top-of-resume as summary zone until first heading resets it
    )
    in_skills = False
    in_certifications = False

    for index, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue
        style = getattr(paragraph, "style", None)
        style_name = getattr(style, "name", "") or ""
        style_lower = style_name.lower()

        is_heading = (
            any(heading in style_lower for heading in heading_styles)
            or _is_section_heading_text(text, date_pattern, section_heading_keywords)
            or style_lower in ("title", "subtitle")
        )
        word_count = len(text.split())

        # Guard: some resume templates style the summary body text with a heading/subtitle
        # style.  When we're already inside a summary section (in_summary=True) and the
        # paragraph is clearly body text (many words, mixed case), DON'T promote it to
        # "heading" — keep in_summary active so it gets the correct "summary" type.
        if is_heading and in_summary and word_count >= 5 and not text.isupper():
            is_heading = False

        if is_heading:
            current_section = text
            lower_section = text.lower()
            in_summary = any(
                kw in lower_section
                for kw in (
                    "summary",
                    "objective",
                    "profile",
                    "about",
                    "career objective",
                    "personal statement",
                )
            )
            in_skills = any(
                keyword in lower_section for keyword in skills_section_names
            )
            in_certifications = any(
                keyword in lower_section for keyword in certification_section_names
            )
            # certifications section must not be treated as skills
            if in_certifications:
                in_skills = False
            result.append(
                {
                    "index": index,
                    "text": text,
                    "style": style_name,
                    "section": current_section,
                    "is_heading": True,
                    "paragraph_type": "heading",
                }
            )
            continue

        paragraph_type = _classify_paragraph_text(
            text,
            in_summary,
            in_skills,
            in_certifications,
            date_pattern,
        )

        result.append(
            {
                "index": index,
                "text": text,
                "style": style_name,
                "section": current_section,
                "is_heading": False,
                "paragraph_type": paragraph_type,
            }
        )

    # Debug: print classification summary so you can verify the summary is tagged correctly
    for p in result:
        if p["paragraph_type"] in ("summary", "certification") or p.get("is_heading"):
            print(
                f"[CLASSIFY] type={p['paragraph_type']!r:15} | {repr(p['text'][:70])}"
            )

    return result


def apply_replacements(original_path, output_path, replacements):
    """
    Apply text replacements to a docx while preserving all formatting.

    Matching strategy (in order):
    1. Exact match after normalize_text()
    2. Fuzzy match (ratio >= 0.85) for cases where LLM slightly altered the original text

    Writing strategy:
    - Preserve the formatting (bold, italic, font, size) of the first run
    - Set first run to full new text, clear all subsequent runs
    - This keeps bullet formatting, indentation, and style intact
    """
    shutil.copy2(original_path, output_path)
    document = Document(output_path)

    replacement_map = _build_replacement_map(replacements)

    if not replacement_map:
        document.save(output_path)
        return

    applied, matched = _apply_replacement_map(
        document, replacement_map, preserve_bold_split=True
    )

    print(f"[INFO] Applied {applied}/{len(replacement_map)} replacements")

    for key in replacement_map:
        if key not in matched:
            print(f"[UNMATCHED] {repr(key[:80])}")

    document.save(output_path)


def apply_title_changes(src_path, dst_path, changes):
    """
    Apply explicit role-title changes chosen by the user to an already-generated docx.
    `changes` is a list of {"original_title": "...", "new_title": "..."} dicts.
    Unlike apply_replacements, this bypasses the paragraph-type filter and operates
    directly on title paragraphs.
    """
    shutil.copy2(src_path, dst_path)
    document = Document(dst_path)

    change_map = {}
    for c in changes:
        old = normalize_text(c.get("original_title", ""))
        new = c.get("new_title", "").strip()
        if old and new:
            change_map[old] = new

    if not change_map:
        document.save(dst_path)
        return

    applied = _apply_replacement_map(document, change_map, preserve_bold_split=False)

    print(f"[INFO] Title changes applied: {applied}/{len(change_map)}")
    document.save(dst_path)


def resume_info_data():
    if not os.path.exists(DEFAULT_RESUME):
        return {"exists": False}

    text = extract_resume_text(DEFAULT_RESUME)
    document = Document(DEFAULT_RESUME)
    paragraphs = [
        paragraph.text.strip()
        for paragraph in document.paragraphs
        if paragraph.text.strip()
    ]
    original_filename = "base_resume.docx"
    if os.path.exists(ORIGINAL_RESUME_INFO):
        with open(ORIGINAL_RESUME_INFO, "r", encoding="utf-8") as file_handle:
            original_filename = file_handle.read().strip()

    return {
        "exists": True,
        "paragraphs": len(paragraphs),
        "words": len(text.split()),
        "filename": "base_resume.docx",
        "original_filename": original_filename,
        "paragraphs_text": paragraphs,
    }


def enrich_with_sections(replacements, resume_paras):
    """Add section context to each replacement dict, including role/project title."""
    current_section = ""
    current_title = ""
    paragraph_lookup = {}
    for paragraph in resume_paras:
        if paragraph.get("is_heading"):
            current_section = paragraph.get("section", "")
            current_title = ""
            continue
        if paragraph.get("paragraph_type") == "title":
            current_title = paragraph["text"]
        key = normalize_text(paragraph["text"])
        label = (
            f"{current_section} > {current_title}" if current_title else current_section
        )
        paragraph_lookup[key] = label

    enriched = []
    for replacement in replacements:
        if not isinstance(replacement, dict):
            continue
        original_norm = normalize_text(replacement.get("original", ""))
        section = _best_section_label(original_norm, paragraph_lookup)
        word_count = len(replacement.get("replacement", "").split())
        enriched.append({**replacement, "section": section, "word_count": word_count})
    return enriched


def extract_role_suggestion(items):  # NOSONAR
    """
    Pop the optional role-suggestion sentinel from the replacements list.
    Supports two formats emitted by the LLM:
      New: {"role_suggestions": [{"original_title": "...", "suggested_title": "..."}, ...]}
      Legacy: {"role_suggestion": "single title string"}
    Returns (list_of_{original_title,suggested_title}_dicts_or_None, remaining_items_list).
    """
    role_suggestions = None
    normal_items = []
    for item in items:
        if not isinstance(item, dict):
            continue
        if "original" in item or "replacement" in item:
            normal_items.append(item)
            continue
        # New array format
        if "role_suggestions" in item:
            raw = item["role_suggestions"]
            if isinstance(raw, list):
                parsed = []
                for s in raw:
                    if isinstance(s, dict) and "suggested_title" in s:
                        parsed.append(
                            {
                                "original_title": str(
                                    s.get("original_title", "")
                                ).strip(),
                                "suggested_title": str(s["suggested_title"]).strip(),
                            }
                        )
                if parsed:
                    role_suggestions = parsed
            continue
        # Legacy single-string format
        if "role_suggestion" in item:
            val = str(item["role_suggestion"]).strip()
            if val:
                role_suggestions = [{"original_title": "", "suggested_title": val}]
    return role_suggestions, normal_items
